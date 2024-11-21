from docx import Document
from docx.shared import Pt  
from io import BytesIO
import base64

def generate_docx(result):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Candara Light'
    font.size = Pt(12)
    doc.add_heading('Ghessebot ebooks', 0)
    if not isinstance(result, str):
        result = str(result) 
    doc.add_paragraph(result)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# download link
def download_docx(bio, filename):
    download_icon = ":inbox_tray:"  # emojis: https://www.webfx.com/tools/emoji-cheat-sheet/
    b64 = base64.b64encode(bio.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download your book {download_icon}</a>'